#!/usr/bin/env python3
"""通用工具函数

包含一些通用的工具函数
"""

import os
import sys
import time
import traceback
from datetime import datetime
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(os.path.dirname(__file__), 'text_capture.log'), encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


def get_timestamp(format='%Y-%m-%d %H:%M:%S'):
    """获取当前时间戳"""
    return datetime.now().strftime(format)


def get_file_size(file_path):
    """获取文件大小"""
    try:
        if os.path.exists(file_path):
            return os.path.getsize(file_path)
        return 0
    except Exception as e:
        logger.error(f"获取文件大小失败 {file_path}: {e}")
        return 0


def ensure_dir_exists(dir_path):
    """确保目录存在"""
    try:
        if not os.path.exists(dir_path):
            os.makedirs(dir_path)
            logger.info(f"目录已创建: {dir_path}")
        return True
    except Exception as e:
        logger.error(f"创建目录失败 {dir_path}: {e}")
        return False


def get_app_data_dir():
    """获取应用程序数据目录"""
    try:
        app_data_dir = os.path.join(os.getenv('APPDATA', os.path.expanduser('~')), 'TextCaptureTool')
        ensure_dir_exists(app_data_dir)
        return app_data_dir
    except Exception as e:
        logger.error(f"获取应用程序数据目录失败: {e}")
        return os.path.expanduser('~')


def get_log_file_path():
    """获取日志文件路径"""
    try:
        app_data_dir = get_app_data_dir()
        return os.path.join(app_data_dir, 'text_capture.log')
    except Exception as e:
        logger.error(f"获取日志文件路径失败: {e}")
        return os.path.join(os.path.expanduser('~'), 'text_capture.log')


def clean_old_logs(days=30):
    """清理旧日志文件"""
    try:
        log_file = get_log_file_path()
        if os.path.exists(log_file):
            file_mod_time = os.path.getmtime(log_file)
            current_time = time.time()
            
            if (current_time - file_mod_time) > days * 24 * 60 * 60:
                os.remove(log_file)
                logger.info(f"旧日志文件已清理: {log_file}")
    except Exception as e:
        logger.error(f"清理旧日志文件失败: {e}")


def validate_file_path(file_path, extension=None):
    """验证文件路径是否有效"""
    try:
        if not file_path:
            return False, "文件路径不能为空"
        
        # 检查文件扩展名
        if extension and not file_path.lower().endswith(extension.lower()):
            return False, f"文件扩展名必须是 {extension}"
        
        # 检查文件所在目录是否存在
        dir_path = os.path.dirname(file_path)
        if dir_path and not os.path.exists(dir_path):
            try:
                os.makedirs(dir_path)
                logger.info(f"文件目录已创建: {dir_path}")
            except Exception as e:
                return False, f"无法创建文件目录: {e}"
        
        return True, "文件路径有效"
    except Exception as e:
        logger.error(f"验证文件路径失败 {file_path}: {e}")
        return False, f"文件路径无效: {e}"


def sanitize_text(text):
    """清理文本内容"""
    try:
        if not isinstance(text, str):
            return ""
        
        # 去除首尾空白字符
        sanitized = text.strip()
        
        # 替换多个换行符为单个换行符
        sanitized = '\n'.join([line.strip() for line in sanitized.split('\n') if line.strip()])
        
        # 替换多个空格为单个空格
        sanitized = ' '.join(sanitized.split())
        
        return sanitized
    except Exception as e:
        logger.error(f"清理文本失败: {e}")
        return ""


def truncate_text(text, max_length=1000):
    """截断文本到指定长度"""
    try:
        if not isinstance(text, str):
            return ""
        
        if len(text) <= max_length:
            return text
        
        return text[:max_length] + "..."
    except Exception as e:
        logger.error(f"截断文本失败: {e}")
        return ""


def get_process_name(pid):
    """根据进程ID获取进程名"""
    try:
        import psutil
        
        process = psutil.Process(pid)
        return process.name()
    except ImportError as e:
        logger.error(f"psutil模块未安装: {e}")
        return "unknown.exe"
    except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess) as e:
        logger.warning(f"无法获取进程名 PID={pid}: {e}")
        return "unknown.exe"
    except Exception as e:
        logger.error(f"获取进程名失败 PID={pid}: {e}")
        return "unknown.exe"


def get_active_window_process_name():
    """获取当前活动窗口的进程名"""
    try:
        import win32gui
        import win32process
        
        # 获取当前活动窗口
        hwnd = win32gui.GetForegroundWindow()
        
        if hwnd == 0:
            return ""
        
        # 获取进程ID
        _, pid = win32process.GetWindowThreadProcessId(hwnd)
        
        if pid == 0:
            return ""
        
        # 获取进程名
        return get_process_name(pid)
    except ImportError as e:
        logger.error(f"win32gui或win32process模块未安装: {e}")
        return ""
    except Exception as e:
        logger.error(f"获取当前活动窗口进程名失败: {e}")
        return ""


def show_notification(title, message, timeout=5000):
    """显示系统通知"""
    try:
        from win10toast import ToastNotifier
        
        toaster = ToastNotifier()
        toaster.show_toast(title, message, duration=timeout/1000)
        return True
    except ImportError as e:
        logger.error(f"win10toast模块未安装: {e}")
        return False
    except Exception as e:
        logger.error(f"显示系统通知失败: {e}")
        return False


def capture_screen(bbox=None):
    """捕获屏幕截图"""
    try:
        from PIL import ImageGrab
        
        if bbox:
            # 确保bbox是有效的
            x1, y1, x2, y2 = bbox
            if x2 <= x1 or y2 <= y1:
                logger.error(f"无效的截图区域: {bbox}")
                return None
            
            return ImageGrab.grab(bbox=(x1, y1, x2, y2))
        else:
            return ImageGrab.grab()
    except ImportError as e:
        logger.error(f"PIL模块未安装: {e}")
        return None
    except Exception as e:
        logger.error(f"捕获屏幕截图失败: {e}")
        return None





def save_text_to_docx(text, docx_path, source_tag=None):
    """保存文本到Word文档"""
    try:
        from docx import Document
        
        # 验证文件路径
        is_valid, error_msg = validate_file_path(docx_path, extension='.docx')
        if not is_valid:
            logger.error(f"保存文本到Word文档失败: {error_msg}")
            return False
        
        # 清理文本
        sanitized_text = sanitize_text(text)
        
        if not sanitized_text:
            logger.warning("无法保存空文本到Word文档")
            return False
        
        # 打开或创建文档
        try:
            if os.path.exists(docx_path):
                doc = Document(docx_path)
            else:
                doc = Document()
                doc.add_heading('文本捕获记录', 0)
        except Exception as e:
            logger.error(f"打开或创建Word文档失败 {docx_path}: {e}")
            return False
        
        # 添加文本段落
        if source_tag:
            doc.add_paragraph(f"{source_tag} {sanitized_text}")
        else:
            doc.add_paragraph(sanitized_text)
        
        # 保存文档
        try:
            doc.save(docx_path)
            logger.info(f"文本已保存到Word文档: {docx_path}")
            return True
        except Exception as e:
            logger.error(f"保存Word文档失败 {docx_path}: {e}")
            return False
    except ImportError as e:
        logger.error(f"python-docx模块未安装: {e}")
        return False
    except Exception as e:
        logger.error(f"保存文本到Word文档失败: {e}")
        return False


def get_selected_text() -> str:
    """
    获取当前系统中选中的文本（通过模拟Ctrl+C操作）
    
    Returns:
        str: 选中的文本内容，如果获取失败则返回空字符串
    """
    try:
        import pyperclip
        import ctypes
        import time
        from ctypes import wintypes
        
        # 定义Windows API函数
        user32 = ctypes.windll.user32
        
        # 保存当前剪贴板内容，以便稍后恢复
        try:
            original_clipboard = pyperclip.paste()
        except Exception:
            original_clipboard = ""
        
        # 清空剪贴板
        pyperclip.copy("")
        
        # 发送Ctrl+C命令复制选中文本（优化时间间隔）
        user32.keybd_event(0x11, 0, 0, 0)  # Ctrl键按下
        time.sleep(0.02)  # 减少等待时间
        user32.keybd_event(0x43, 0, 0, 0)  # C键按下
        time.sleep(0.1)   # 大幅减少等待时间
        user32.keybd_event(0x43, 0, 2, 0)  # C键释放
        time.sleep(0.02)  # 减少等待时间
        user32.keybd_event(0x11, 0, 2, 0)  # Ctrl键释放
        
        # 等待系统处理复制操作（优化时间）
        time.sleep(0.05)  # 大幅减少等待时间
        
        # 获取剪贴板内容
        copied_text = pyperclip.paste()
        logger.debug(f"从剪贴板获取到文本: {repr(copied_text)}")
        
        # 如果获取到的文本为空，可能是没有选中文本
        if not copied_text or copied_text.strip() == "":
            logger.warning("剪贴板中没有可用的选中文本")
            # 恢复原来的剪贴板内容
            if original_clipboard:
                pyperclip.copy(original_clipboard)
            return ""
        
        # 恢复原来的剪贴板内容
        if original_clipboard:
            pyperclip.copy(original_clipboard)
        
        return sanitize_text(copied_text)
        
    except ImportError:
        logger.error("pyperclip模块未安装")
        return ""
    except Exception as e:
        logger.error(f"获取选中文本时发生错误: {e}", exc_info=True)
        return ""


def format_exception(e):
    """格式化异常信息"""
    try:
        return f"{type(e).__name__}: {e}"
    except Exception as ex:
        return f"格式化异常信息失败: {ex}"


def is_user_active():
    """检测用户是否处于活动状态（通过检测鼠标或键盘活动）"""
    try:
        import ctypes
        from ctypes import wintypes
        
        # 定义Windows API结构体
        class LASTINPUTINFO(ctypes.Structure):
            _fields_ = [
                ("cbSize", wintypes.UINT),
                ("dwTime", wintypes.DWORD)
            ]
        
        # 获取系统最后输入时间
        last_input_info = LASTINPUTINFO()
        last_input_info.cbSize = ctypes.sizeof(LASTINPUTINFO)
        
        if ctypes.windll.user32.GetLastInputInfo(ctypes.byref(last_input_info)):
            # 获取系统启动后的运行时间（毫秒）
            current_time = ctypes.windll.kernel32.GetTickCount()
            
            # 计算空闲时间（秒）
            idle_time = (current_time - last_input_info.dwTime) // 1000
            
            # 如果用户空闲时间超过30秒，认为用户不活跃
            return idle_time < 30
        
        return True  # 如果无法获取信息，默认认为用户活跃
        
    except Exception as e:
        logger.debug(f"检测用户活动状态失败: {e}")
        return True  # 如果出错，默认认为用户活跃


def log_exception(e, message=""):
    """记录异常信息"""
    try:
        if message:
            logger.error(f"{message}: {format_exception(e)}")
        else:
            logger.error(f"异常: {format_exception(e)}")
        
        # 记录完整的堆栈跟踪
        logger.error(f"堆栈跟踪:\n{traceback.format_exc()}")
    except Exception as ex:
        logger.error(f"记录异常信息失败: {format_exception(ex)}")


