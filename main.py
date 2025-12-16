#!/usr/bin/env python3
"""全局文本捕获+自动写入DOCX软件 - 主程序

监听鼠标选中的文本，自动抓取并保存到docx文档中
支持区域OCR识别（右键截图）
抓取/识别的文本按来源打标签
"""

import sys
import os
import time
import threading
from datetime import datetime

from PyQt5.QtWidgets import QApplication, QSystemTrayIcon, QMenu, QAction, QFileDialog, QMessageBox, QWidget, QLabel, QLineEdit, QPushButton, QVBoxLayout, QHBoxLayout, QSpinBox, QStyle, QShortcut
from PyQt5.QtGui import QIcon, QCursor, QPainter, QPen, QBrush, QColor
from PyQt5.QtCore import Qt, QTimer, QPoint, QRect, QThread, pyqtSignal

# 导入自定义模块
import config
import utils



class TextCaptureApp(QApplication):
    """主应用程序类"""
    
    def __init__(self, argv):
        super().__init__(argv)
        
        # 设置应用程序属性，确保托盘图标存在时不退出
        self.setQuitOnLastWindowClosed(False)
        
        # 初始化应用程序设置
        self.settings = {
            'capture_enabled': False,
            'docx_path': config.config.get_docx_path(),
            'last_selected_text': '',
            'last_selection_time': 0,
            'capture_count': 0,
        }
        
        # 初始化超时保护相关变量
        self.capture_start_time = 0
        self.max_capture_time = 300  # 最大捕获时间，默认300秒（5分钟）
        self.max_capture_count = 10000   # 最大捕获次数，默认10000次
        
        # 初始化截图相关变量
        self.is_screenshotting = False
        self.screenshot_start = QPoint()
        self.screenshot_end = QPoint()
        self.screenshot_widget = None
        
        # 初始化文档
        self.init_document()
        
        # 创建系统托盘图标
        self.create_tray_icon()
        
        # 创建捕获线程
        self.capture_thread = None
        
        # 创建捕获定时器（用于单次检查）
        self.capture_timer = QTimer()
        self.capture_timer.timeout.connect(self.check_selection)
        self.capture_timer.setInterval(int(config.config.get_capture_interval() * 1000))
        
        # 创建OCR热键
        self.register_hotkeys()
        
        utils.logger.info("应用程序初始化完成")
        
    def init_document(self):
        """初始化文档"""
        try:
            from docx import Document
            
            # 获取程序所在目录
            app_dir = os.path.dirname(os.path.abspath(__file__))
            # 使用程序所在目录作为默认文档路径
            self.settings['docx_path'] = os.path.join(app_dir, 'text_capture.docx')
            
            if os.path.exists(self.settings['docx_path']):
                self.document = Document(self.settings['docx_path'])
            else:
                self.document = Document()
                self.document.add_heading('文本捕获记录', 0)
                self.document.save(self.settings['docx_path'])
            utils.logger.info(f"文档已初始化：{self.settings['docx_path']}")
        except Exception as e:
            utils.logger.error(f"初始化文档失败：{e}")
            QMessageBox.critical(None, '错误', f'无法创建或打开文档：\n{self.settings["docx_path"]}\n\n错误：{e}')
            sys.exit(1)
            
    def create_tray_icon(self):
        """创建系统托盘图标"""
        # 创建托盘图标（使用Qt内置的信息图标）
        self.tray_icon = QSystemTrayIcon()
        self.tray_icon.setIcon(self.style().standardIcon(QStyle.SP_MessageBoxInformation))
        self.tray_icon.setToolTip('文本捕获工具')
        
        # 显示托盘图标
        self.tray_icon.show()
        
        # 创建菜单
        tray_menu = QMenu()
        
        # 开始/停止捕获动作（动态切换）
        self.start_stop_action = QAction('开始捕获', self)
        self.start_stop_action.triggered.connect(self.toggle_capture)
        tray_menu.addAction(self.start_stop_action)
        
        tray_menu.addSeparator()
        
        # 移除OCR截图动作，只保留快捷键启动方式
        # ocr_action = QAction('区域OCR识别', self)
        # ocr_action.triggered.connect(self.start_screenshot)
        # tray_menu.addAction(ocr_action)
        
        # 设置动作
        settings_action = QAction('设置', self)
        settings_action.triggered.connect(self.show_settings)
        tray_menu.addAction(settings_action)
        
        # 查看文档动作
        view_doc_action = QAction('查看文档', self)
        view_doc_action.triggered.connect(self.open_document)
        tray_menu.addAction(view_doc_action)
        
        # 关于动作
        about_action = QAction('关于', self)
        about_action.triggered.connect(self.show_about)
        tray_menu.addAction(about_action)
        
        # 退出动作
        exit_action = QAction('退出', self)
        exit_action.triggered.connect(self.quit)
        tray_menu.addAction(exit_action)
        
        # 设置托盘菜单
        self.tray_icon.setContextMenu(tray_menu)
        
        # 显示托盘图标
        self.tray_icon.show()
        
    def create_icon(self):
        """创建简单的图标"""
        from io import BytesIO
        from PIL import Image, ImageDraw, ImageFont
        from PyQt5.QtGui import QPixmap, QByteArray
        from PyQt5.QtCore import QBuffer
        
        # 创建16x16像素的图标
        img = Image.new('RGB', (16, 16), color='blue')
        d = ImageDraw.Draw(img)
        
        try:
            # 使用系统默认字体
            font = ImageFont.truetype("arial.ttf", 8)
        except:
            font = ImageFont.load_default()
            
        d.text((1, 1), "T", fill='white', font=font)
        
        # 保存到内存
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        
        # 转换为QPixmap
        pixmap = QPixmap()
        pixmap.loadFromData(buffer.getvalue())
        
        return pixmap
        
    def toggle_capture(self):
        """切换捕获状态"""
        if self.settings['capture_enabled']:
            # 停止捕获
            self.settings['capture_enabled'] = False
            self.start_stop_action.setText('开始捕获')
            self.tray_icon.setIcon(self.style().standardIcon(QStyle.SP_MessageBoxInformation))
            
            # 停止捕获线程
            if self.capture_thread and self.capture_thread.isRunning():
                self.capture_thread.stop()
                self.capture_thread.wait()
                self.capture_thread = None
            
            # 停止定时器
            self.capture_timer.stop()
            
            # 弹出保存文档对话框让用户选择路径
            self.show_save_document_dialog()
            
            utils.logger.info("文本捕获已停止")
        else:
            # 开始捕获
            self.settings['capture_enabled'] = True
            self.start_stop_action.setText('停止捕获')
            self.tray_icon.setIcon(self.style().standardIcon(QStyle.SP_MessageBoxCritical))
            
            # 重置捕获计数
            self.settings['capture_count'] = 0
            self.settings['last_selected_text'] = ''
            self.settings['last_selection_time'] = 0
            
            # 每次开始捕获时创建新的文档
            try:
                from docx import Document
                from datetime import datetime
                
                # 生成带时间戳的新文档文件名
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                app_dir = os.path.dirname(os.path.abspath(__file__))
                new_docx_path = os.path.join(app_dir, f'text_capture_{timestamp}.docx')
                
                # 创建新文档
                self.document = Document()
                self.document.add_heading('文本捕获记录', 0)
                
                # 保存新文档并更新设置
                self.document.save(new_docx_path)
                self.settings['docx_path'] = new_docx_path
                
                utils.logger.info(f"创建新文档：{new_docx_path}")
                
            except Exception as e:
                utils.logger.error(f"创建新文档失败：{e}")
                self.tray_icon.showMessage('文本捕获工具', f'无法创建新文档：{e}', QSystemTrayIcon.Critical, 3000)
                return
            
            # 创建并启动捕获线程
            self.capture_thread = self.CaptureThread(self.max_capture_time, self.max_capture_count)
            self.capture_thread.text_captured.connect(self.handle_text_captured)
            self.capture_thread.start()
            
            self.tray_icon.showMessage('文本捕获工具', f'已开始文本捕获，将持续{self.max_capture_time//60}分钟或捕获{self.max_capture_count}次后自动停止', QSystemTrayIcon.Information, 3000)
            utils.logger.info(f"文本捕获已开始（持续模式），超时保护：{self.max_capture_time}秒/{self.max_capture_count}次")
    
    def show_save_document_dialog(self):
        """显示保存文档对话框让用户选择保存路径"""
        try:
            # 如果没有捕获到任何文本，直接返回
            if self.settings['capture_count'] == 0:
                self.tray_icon.showMessage('文本捕获工具', '没有捕获到任何文本需要保存', QSystemTrayIcon.Information, 3000)
                return
            
            # 弹出保存文件对话框
            file_path, _ = QFileDialog.getSaveFileName(
                None, 
                '保存文本捕获记录', 
                os.path.join(os.path.expanduser('~'), 'Documents', f'text_capture_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'), 
                'Word文档 (*.docx)'
            )
            
            if file_path:
                # 如果用户选择了新的保存路径，更新设置并保存文档
                if self.settings['docx_path'] != file_path:
                    # 创建新文档
                    from docx import Document
                    new_doc = Document()
                    new_doc.add_heading('文本捕获记录', 0)
                    
                    # 将原文档内容复制到新文档
                    if os.path.exists(self.settings['docx_path']):
                        original_doc = Document(self.settings['docx_path'])
                        for para in original_doc.paragraphs[1:]:  # 跳过标题
                            new_doc.add_paragraph(para.text)
                    
                    new_doc.save(file_path)
                    self.settings['docx_path'] = file_path
                    
                    # 更新文档引用
                    self.document = new_doc
                    
                    self.tray_icon.showMessage('文本捕获工具', f'文档已保存到：{os.path.basename(file_path)}', QSystemTrayIcon.Information, 3000)
                    utils.logger.info(f"文档已保存到：{file_path}")
                    
                    # 重置捕获计数
                    self.settings['capture_count'] = 0
                else:
                    # 路径未改变，直接提示文档已存在
                    self.tray_icon.showMessage('文本捕获工具', f'文档已存在：{os.path.basename(file_path)}', QSystemTrayIcon.Information, 3000)
                    
        except Exception as e:
            utils.logger.error(f"保存文档失败：{e}")
            self.tray_icon.showMessage('文本捕获工具', f'保存文档失败：{e}', QSystemTrayIcon.Critical, 3000)
            
    class CaptureThread(QThread):
        """捕获线程类 - 在后台持续捕获文本"""
        text_captured = pyqtSignal(str)  # 定义信号，用于发送捕获到的文本
        
        def __init__(self, max_capture_time=300, max_capture_count=10000):
            super().__init__()
            self.is_running = True
            self.timeout = max_capture_time  # 5分钟超时
            self.max_times = max_capture_count  # 最多捕获10000次
            self.capture_count = 0
            self.start_time = 0
            self.last_capture_time = 0
            self.consecutive_empty_count = 0  # 连续空捕获计数
        
        def run(self):
            """线程运行函数"""
            self.start_time = time.time()
            
            # 循环捕获，但不阻塞主线程
            while self.is_running and (time.time() - self.start_time < self.timeout) and (self.capture_count < self.max_times):
                try:
                    current_time = time.time()
                    
                    # 智能捕获间隔控制
                    if not utils.is_user_active():
                        # 用户不活跃时，大幅降低捕获频率
                        sleep_time = 5000  # 5秒
                    elif self.consecutive_empty_count > 5:
                        # 如果连续多次没有捕获到文本，降低捕获频率
                        sleep_time = 2000  # 2秒
                    elif self.consecutive_empty_count > 2:
                        # 如果连续几次没有捕获到文本，适当降低频率
                        sleep_time = 1000  # 1秒
                    else:
                        # 正常捕获频率
                        sleep_time = 500  # 0.5秒
                    
                    selected_text = utils.get_selected_text()  # 调用捕获函数
                    
                    if selected_text:
                        self.text_captured.emit(selected_text)  # 发信号给主线程
                        self.capture_count += 1
                        self.consecutive_empty_count = 0  # 重置空捕获计数
                        # 捕获到文本后等待一段时间，避免过于频繁
                        self.msleep(1000)  # 捕获成功后等待1秒
                    else:
                        self.consecutive_empty_count += 1
                        # 使用QThread的sleep方法，避免阻塞主线程
                        self.msleep(sleep_time)
                    
                except Exception as e:
                    utils.logger.error(f"捕获文本时出错：{e}")
                    import traceback
                    utils.logger.error(traceback.format_exc())
                    self.msleep(1000)  # 出错后等待1秒
        
        def stop(self):
            """停止线程"""
            self.is_running = False
    
    def handle_text_captured(self, selected_text):
        """处理捕获到的文本"""
        try:
            current_time = time.time()
            
            # 快速检查：如果文本与上次捕获的完全相同，直接跳过
            if selected_text == self.settings['last_selected_text']:
                utils.logger.debug("检测到重复文本，已跳过处理")
                return
            
            # 添加详细日志
            utils.logger.debug(f"检测到选中文本：{utils.truncate_text(selected_text, 100)}")
            utils.logger.debug(f"文本长度：{len(selected_text)}")
            
            # 检查文本是否有效
            min_length = config.config.get_min_text_length()
            max_length = config.config.get_max_text_length()
            
            if selected_text and \
               len(selected_text) >= min_length and \
               len(selected_text) <= max_length:
                
                # 检查时间间隔，避免重复捕获
                if current_time - self.settings['last_selection_time'] > 2:
                    # 获取文本来源
                    source_tag = self.get_text_source()
                    
                    # 保存文本
                    self.save_text(selected_text, source_tag)
                    
                    # 更新状态
                    self.settings['last_selected_text'] = selected_text
                    self.settings['last_selection_time'] = current_time
                    self.settings['capture_count'] += 1
                    
                    utils.logger.info(f"捕获到文本：{utils.truncate_text(selected_text, 50)}")
                    utils.logger.info(f"来源：{source_tag}")
                    utils.logger.info(f"总捕获数：{self.settings['capture_count']}")
                else:
                    utils.logger.debug("时间间隔过短，跳过重复捕获")
            else:
                utils.logger.debug("文本无效或长度不符合要求，已跳过")
                    
        except Exception as e:
            # 记录详细错误日志
            utils.logger.error(f"处理捕获到的文本时发生错误：{e}")
            import traceback
            utils.logger.error(traceback.format_exc())
    
    def check_selection(self):
        """检查当前选中的文本（单次检查）"""
        try:
            # 只有在捕获未启用时才进行单次检查
            if not self.settings['capture_enabled']:
                selected_text = utils.get_selected_text()
                if selected_text:
                    self.handle_text_captured(selected_text)
                    
        except Exception as e:
            # 记录详细错误日志
            utils.logger.error(f"检查选中文本时发生错误：{e}")
            import traceback
            utils.logger.error(traceback.format_exc())
            
    def get_text_source(self):
        """获取当前活动窗口的进程名，判断文本来源"""
        try:
            process_name = utils.get_active_window_process_name()
            
            if process_name:
                # 根据进程名返回对应的标签
                return config.config.get_text_source_tag(process_name)
            else:
                return '[未知来源]'
                
        except Exception as e:
            # 如果无法获取来源信息，返回默认标签
            utils.logger.debug(f"获取文本来源时发生错误：{e}")
            return '[未知来源]'
            
    def browse_docx_path(self):
        """浏览选择DOCX文档路径"""
        try:
            file_path, _ = QFileDialog.getSaveFileName(None, '选择保存文档', self.docx_path_edit.text(), 'Word文档 (*.docx)')
            
            if file_path:
                self.docx_path_edit.setText(file_path)
                
        except Exception as e:
            utils.logger.error(f"浏览DOCX文档路径失败：{e}")
            self.tray_icon.showMessage('文本捕获工具', f'浏览DOCX文档路径失败：{e}', QSystemTrayIcon.Critical, 3000)
            

            
    def save_settings(self):
        """保存设置"""
        try:
            # 保存DOCX文档路径
            docx_path = self.docx_path_edit.text()
            if docx_path:
                self.settings['docx_path'] = docx_path
                
                # 重新初始化文档
                self.init_document()
                

                
            # 保存捕获间隔（转换为秒，因为配置文件中使用秒为单位）
            capture_interval_ms = self.interval_spin.value()
            capture_interval_sec = capture_interval_ms / 1000.0
            config.config.set_capture_interval(capture_interval_sec)
            
            # 保存最小文本长度
            min_text_length = self.min_length_spin.value()
            config.config.set_min_text_length(min_text_length)
            
            # 保存最大文本长度
            max_text_length = self.max_length_spin.value()
            config.config.set_max_text_length(max_text_length)
            
            # 保存配置到文件
            config.config.save_config()
            
            # 更新捕获定时器间隔（使用毫秒）
            if self.capture_timer:
                self.capture_timer.stop()
                self.capture_timer.start(capture_interval_ms)
                
            utils.logger.info("设置已保存")
            self.tray_icon.showMessage('文本捕获工具', '设置已保存', QSystemTrayIcon.Information, 3000)
            
            # 关闭设置窗口
            if hasattr(self, 'settings_window') and self.settings_window:
                self.settings_window.close()
                self.settings_window = None
                
        except Exception as e:
            utils.logger.error(f"保存设置失败：{e}")
            self.tray_icon.showMessage('文本捕获工具', f'保存设置失败：{e}', QSystemTrayIcon.Critical, 3000)
            
    def save_text(self, text, source_tag='[未知来源]'):
        """保存文本到DOCX文档"""
        try:
            # 添加详细日志
            utils.logger.debug(f"准备保存文本：{utils.truncate_text(text, 100)}")
            utils.logger.debug(f"文本来源：{source_tag}")
            utils.logger.debug(f"保存路径：{self.settings['docx_path']}")
            
            # 清理文本
            text = utils.sanitize_text(text)
            
            utils.logger.debug(f"清理后的文本：{utils.truncate_text(text, 100)}")
            
            if not text:
                utils.logger.debug("尝试保存空文本，已跳过")
                return
                
            # 使用utils.py中的方法保存文本
            success = utils.save_text_to_docx(text, self.settings['docx_path'], source_tag)
            
            if not success:
                raise Exception("保存文本到Word文档失败")
            else:
                utils.logger.info("文本保存成功")
            
        except Exception as e:
            utils.logger.error(f"保存文本失败：{e}")
            import traceback
            utils.logger.error(traceback.format_exc())
            self.tray_icon.showMessage('文本捕获工具', f'保存文本失败：{e}', QSystemTrayIcon.Critical, 3000)
            
    def start_screenshot(self):
        """开始截图"""
        self.is_screenshotting = False  # 初始状态为未开始选择
        self.has_started_selection = False  # 是否已经开始选择
        self.screenshot_start = None
        self.screenshot_end = None
        
        # 创建截图区域选择窗口
        self.screenshot_widget = QWidget()
        self.screenshot_widget.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint)
        
        # 设置窗口为完全透明（不绘制背景）
        self.screenshot_widget.setAttribute(Qt.WA_TranslucentBackground)
        
        # 获取整个屏幕的大小（适合截图全屏）
        screen_rect = QApplication.desktop().screenGeometry()
        self.screenshot_widget.setGeometry(screen_rect)
        
        # 添加绘制事件
        self.screenshot_widget.paintEvent = self.on_screenshot_paint
        
        # 设置鼠标跟踪，确保能捕获鼠标移动事件
        self.screenshot_widget.setMouseTracking(True)
        
        # 确保窗口获得焦点
        self.screenshot_widget.raise_()
        self.screenshot_widget.activateWindow()
        
        self.screenshot_widget.show()
        
        # 捕获全局鼠标事件
        self.screenshot_widget.grabMouse()
        
        # 监听鼠标按下、移动和释放事件
        self.screenshot_widget.mousePressEvent = self.on_screenshot_mouse_press
        self.screenshot_widget.mouseMoveEvent = self.on_screenshot_mouse_move
        self.screenshot_widget.mouseReleaseEvent = self.on_screenshot_mouse_release
        
        utils.logger.info("开始截图，等待第一次点击设置起始点")
        
    def on_screenshot_paint(self, event):
        """截图窗口绘制事件"""
        painter = QPainter(self.screenshot_widget)
        
        # 禁用抗锯齿以提高性能
        # painter.setRenderHint(QPainter.Antialiasing, False)
        
        # 不绘制任何背景，保持窗口完全透明，让用户能看到屏幕内容
        
        # 如果已经开始选择，绘制十字准星和选择区域
        if self.is_screenshotting and self.has_started_selection:
            # 绘制十字准星（使用简单线条，提高性能）
            painter.setPen(QPen(QColor(255, 0, 0), 1))
            
            # 水平线
            painter.drawLine(0, self.screenshot_end.y(), 
                            self.screenshot_widget.width(), self.screenshot_end.y())
            # 垂直线
            painter.drawLine(self.screenshot_end.x(), 0,
                            self.screenshot_end.x(), self.screenshot_widget.height())
            
            # 绘制选择区域（只要起始点和结束点不同就绘制）
            if self.screenshot_start != self.screenshot_end:
                x1 = min(self.screenshot_start.x(), self.screenshot_end.x())
                y1 = min(self.screenshot_start.y(), self.screenshot_end.y())
                x2 = max(self.screenshot_start.x(), self.screenshot_end.x())
                y2 = max(self.screenshot_start.y(), self.screenshot_end.y())
                
                # 绘制选择区域边框（使用简单边框）
                painter.setPen(QPen(QColor(255, 0, 0), 2))
                painter.drawRect(x1, y1, x2 - x1, y2 - y1)
                
                # 绘制选择区域内部（简化透明度）
                painter.fillRect(x1, y1, x2 - x1, y2 - y1, QColor(0, 0, 0, 60))
                
                # 只在区域很大时才显示尺寸信息，减少文本绘制开销
                if (x2 - x1) > 200 and (y2 - y1) > 100:
                    painter.setPen(QColor(255, 0, 0))
                    painter.setBrush(QColor(255, 255, 255, 200))
                    painter.drawRect(x1, y1 - 25, 100, 20)
                    painter.setPen(QColor(0, 0, 0))
                    painter.drawText(x1 + 5, y1 - 10, f"{x2 - x1} x {y2 - y1}")
        
        painter.end()
        
    def on_screenshot_mouse_press(self, event):
        """截图时鼠标按下事件"""
        if event.button() == Qt.LeftButton:
            if not self.has_started_selection:
                # 第一次点击：设置起始点（使用窗口相对坐标）
                self.screenshot_start = event.pos()
                self.screenshot_end = self.screenshot_start
                self.has_started_selection = True
                self.is_screenshotting = True
                utils.logger.info(f"设置选择起始点：{self.screenshot_start}")
            else:
                # 第二次点击：结束选择
                utils.logger.info("结束选择区域，执行OCR识别")
                
                # 计算截图区域（使用全局坐标进行截图）
                if self.screenshot_start and self.screenshot_end:
                    # 获取窗口在屏幕上的位置
                    window_pos = self.screenshot_widget.mapToGlobal(QPoint(0, 0))
                    
                    # 将窗口相对坐标转换为全局坐标
                    global_start = self.screenshot_widget.mapToGlobal(self.screenshot_start)
                    global_end = self.screenshot_widget.mapToGlobal(self.screenshot_end)
                    
                    x1 = min(global_start.x(), global_end.x())
                    y1 = min(global_start.y(), global_end.y())
                    x2 = max(global_start.x(), global_end.x())
                    y2 = max(global_start.y(), global_end.y())
                    
                    # 确保区域有效
                    if x2 - x1 > 10 and y2 - y1 > 10:
                        # 截图
                        screenshot = utils.capture_screen(bbox=(x1, y1, x2, y2))
                        
                        # OCR识别
                        if screenshot is not None:
                            self.perform_ocr(screenshot)
                        else:
                            utils.logger.error("截图失败，无法进行OCR识别")
                    else:
                        utils.logger.info("截图区域太小，取消截图")
                
                # 完全退出OCR识别模式
                self.exit_screenshot_mode()
            
            # 重绘窗口
            self.screenshot_widget.update()
    
    def exit_screenshot_mode(self):
        """完全退出截图模式"""
        # 隐藏截图窗口
        if self.screenshot_widget:
            self.screenshot_widget.hide()
        
        # 重置所有状态
        self.is_screenshotting = False
        self.has_started_selection = False
        self.screenshot_start = None
        self.screenshot_end = None
        
        # 释放鼠标捕获
        if self.screenshot_widget:
            self.screenshot_widget.releaseMouse()
            self.screenshot_widget.releaseKeyboard()
        
        utils.logger.info("已退出OCR识别模式")
    
    def on_screenshot_mouse_move(self, event):
        """截图时鼠标移动事件"""
        if self.is_screenshotting and self.has_started_selection:
            # 使用窗口相对坐标，确保绘制位置正确
            new_pos = event.pos()
            
            # 只有当鼠标位置有较大变化时才更新，大幅减少重绘频率
            if (self.screenshot_end is None or 
                abs(new_pos.x() - self.screenshot_end.x()) > 3 or 
                abs(new_pos.y() - self.screenshot_end.y()) > 3):
                
                self.screenshot_end = new_pos
                # 使用update()而不是repaint()，让Qt优化重绘时机
                self.screenshot_widget.update()
        
    def on_screenshot_mouse_release(self, event):
        """截图时鼠标释放事件"""
        # 现在通过鼠标按下事件处理选择逻辑
        pass
        

            
    def register_hotkeys(self):
        """注册快捷键"""
        try:
            from PyQt5.QtGui import QKeySequence
            
            # 创建一个隐藏的窗口作为快捷键的父部件
            self.hotkey_window = QWidget()
            self.hotkey_window.setWindowFlags(Qt.Tool | Qt.FramelessWindowHint)
            self.hotkey_window.setVisible(False)
            
            # 注释掉OCR快捷键注册，取消OCR功能
            # shortcut = QShortcut(QKeySequence('Ctrl+Alt+O'), self.hotkey_window)
            # shortcut.activated.connect(self.start_screenshot)
            
            utils.logger.info("已取消OCR快捷键注册")
            
        except Exception as e:
            utils.logger.error(f"注册快捷键失败：{e}")
            
    def show_settings(self):
        """显示设置界面"""
        try:
            # 创建设置窗口
            settings_window = QWidget()
            settings_window.setWindowTitle('文本捕获工具 - 设置')
            settings_window.setGeometry(200, 200, 500, 400)
            settings_window.setWindowFlags(Qt.WindowStaysOnTopHint)
            
            # 创建布局
            layout = QVBoxLayout()
            
            # DOCX文档路径设置
            docx_layout = QHBoxLayout()
            docx_label = QLabel('DOCX文档路径：')
            self.docx_path_edit = QLineEdit(self.settings['docx_path'])
            self.docx_path_edit.setReadOnly(True)
            docx_browse_btn = QPushButton('浏览')
            docx_browse_btn.clicked.connect(self.browse_docx_path)
            
            docx_layout.addWidget(docx_label)
            docx_layout.addWidget(self.docx_path_edit)
            docx_layout.addWidget(docx_browse_btn)
            layout.addLayout(docx_layout)
            

            
            # 捕获间隔设置
            interval_layout = QHBoxLayout()
            interval_label = QLabel('捕获间隔（毫秒）：')
            self.interval_spin = QSpinBox()
            self.interval_spin.setRange(500, 5000)
            # 将浮点数秒转换为整数毫秒
            interval_ms = int(config.config.get_capture_interval() * 1000)
            self.interval_spin.setValue(interval_ms)
            
            interval_layout.addWidget(interval_label)
            interval_layout.addWidget(self.interval_spin)
            layout.addLayout(interval_layout)
            
            # 最小文本长度设置
            min_length_layout = QHBoxLayout()
            min_length_label = QLabel('最小文本长度：')
            self.min_length_spin = QSpinBox()
            self.min_length_spin.setRange(1, 100)
            self.min_length_spin.setValue(config.config.get_min_text_length())
            
            min_length_layout.addWidget(min_length_label)
            min_length_layout.addWidget(self.min_length_spin)
            layout.addLayout(min_length_layout)
            
            # 最大文本长度设置
            max_length_layout = QHBoxLayout()
            max_length_label = QLabel('最大文本长度：')
            self.max_length_spin = QSpinBox()
            self.max_length_spin.setRange(100, 10000)
            self.max_length_spin.setValue(config.config.get_max_text_length())
            
            max_length_layout.addWidget(max_length_label)
            max_length_layout.addWidget(self.max_length_spin)
            layout.addLayout(max_length_layout)
            
            # 按钮布局
            button_layout = QHBoxLayout()
            save_btn = QPushButton('保存设置')
            save_btn.clicked.connect(self.save_settings)
            cancel_btn = QPushButton('取消')
            cancel_btn.clicked.connect(settings_window.close)
            
            button_layout.addWidget(save_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            # 设置布局
            settings_window.setLayout(layout)
            
            # 显示窗口
            settings_window.show()
            
            # 保存窗口引用，防止被垃圾回收
            self.settings_window = settings_window
            
        except Exception as e:
            utils.logger.error(f"显示设置界面失败：{e}")
            self.tray_icon.showMessage('文本捕获工具', f'显示设置界面失败：{e}', QSystemTrayIcon.Critical, 3000)
            
    def open_document(self):
        """打开文档"""
        try:
            if os.path.exists(self.settings['docx_path']):
                os.startfile(self.settings['docx_path'])
                utils.logger.info(f"文档已打开：{self.settings['docx_path']}")
            else:
                utils.logger.warning("打开文档失败：文档不存在")
                self.tray_icon.showMessage('文本捕获工具', '文档不存在', QSystemTrayIcon.Warning, 3000)
                
        except Exception as e:
            utils.logger.error(f"打开文档失败：{e}")
            self.tray_icon.showMessage('文本捕获工具', f'打开文档失败：{e}', QSystemTrayIcon.Critical, 3000)
            
    def show_about(self):
        """显示关于对话框"""
        try:
            about_text = "文本捕获工具 v1.0\n\n" \
                        "一个能够自动捕获并记录全局文本选择内容的工具。\n" \
                        "支持快捷键操作、区域OCR识别、文本自动保存到Word文档等功能。\n\n" \
                        "开发者：出久君\n" \
                        "联系邮箱：2744314855@qq.com\n\n" \
                        "版权所有 © 2025" \
                        
            QMessageBox.about(None, '关于文本捕获工具', about_text)
            utils.logger.info("显示关于对话框")
            
        except Exception as e:
            utils.logger.error(f"显示关于对话框失败：{e}")
            self.tray_icon.showMessage('文本捕获工具', f'显示关于对话框失败：{e}', QSystemTrayIcon.Critical, 3000)
            
    def create_icon(self, active=False):
        """创建简单的图标"""
        from io import BytesIO
        from PIL import Image, ImageDraw, ImageFont
        
        # 创建16x16像素的图标
        img = Image.new('RGB', (16, 16), color='blue' if active else 'gray')
        d = ImageDraw.Draw(img)
        
        try:
            # 使用系统默认字体
            font = ImageFont.truetype("arial.ttf", 8)
        except:
            font = ImageFont.load_default()
            
        d.text((1, 1), "T", fill='white', font=font)
        
        # 保存到内存
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        
        return buffer
        
    def run(self):
        """运行应用程序"""
        print("应用程序已启动")
        utils.logger.info("应用程序主循环已启动")
        return self.exec_()
        

def main():
    """主函数"""
    try:
        # 创建应用程序
        app = TextCaptureApp(sys.argv)
        
        # 运行应用程序
        sys.exit(app.run())
        
    except Exception as e:
        print(f"应用程序启动失败：{e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
        

if __name__ == '__main__':
    main()